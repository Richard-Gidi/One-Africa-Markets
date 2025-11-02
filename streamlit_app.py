# ==============================================================================
# 1. IMPORTS
# ==============================================================================
import os
import re
import html
import json
import hashlib
import datetime as dt
from typing import List, Dict, Tuple, Optional, Any
from urllib.parse import urlparse, urljoin
import math
import time
import logging
from io import BytesIO
from concurrent.futures import ThreadPoolExecutor, as_completed

# --- Core Libs ---
import numpy as np
import pandas as pd
import requests
import streamlit as st
from bs4 import BeautifulSoup
from requests.adapters import HTTPAdapter
from requests.packages.urllib3.util.retry import Retry
import xml.etree.ElementTree as ET

# --- Optional Libs (with fallbacks) ---
try:
    from dotenv import load_dotenv
    load_dotenv()
except Exception:
    pass

try:
    from openai import OpenAI
    OPENAI_OK = True
except Exception:
    OPENAI_OK = False

try:
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity
    HAS_SK = True
except Exception:
    HAS_SK = False

try:
    import nltk
    from nltk.sentiment import SentimentIntensityAnalyzer
    try:
        _ = nltk.data.find("sentiment/vader_lexicon.zip")
    except LookupError:
        nltk.download("vader_lexicon")
    HAS_VADER = True
except Exception:
    HAS_VADER = False

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except Exception:
    HAS_DOCX = False

# ==============================================================================
# 2. CONSTANTS & CONFIGURATION
# ==============================================================================

# --- App Strings ---
APP_NAME = "One Africa Market Pulse"
TAGLINE = "Automated intelligence for cashew, shea, cocoa & allied markets."
QUOTE = "“Ask your data why, until it has nothing else to say.” — Richard Gidi"
FALLBACK_IMG = "https://images.unsplash.com/photo-1519681393784-d120267933ba?q=80&w=1200&auto=format&fit=crop"

# --- Keywords & Sources ---
DEFAULT_KEYWORDS = [
    "cashew", "shea", "shea nut", "cocoa", "palm kernel", "agri", "export", "harvest",
    "shipment", "freight", "logistics", "port", "tariff", "ban", "fx", "currency",
    "cedi", "naira", "inflation", "subsidy", "cooperative", "value-addition", "processing",
    "ghana", "nigeria", "cote d’ivoire", "ivory coast", "benin", "togo", "burkina",
    "west africa", "sahel", "trade policy", "commodity", "price", "market"
]

DEFAULT_SOURCES = {
    "AllAfrica » Agriculture": "https://allafrica.com/tools/headlines/rdf/agriculture/headlines.rdf",
    "AllAfrica » Business": "https://allafrica.com/tools/headlines/rdf/business/headlines.rdf",
    "The Standard » Business": "https://www.standardmedia.co.ke/rss/business.php",
    "The Standard » Agriculture": "https://www.standardmedia.co.ke/rss/agriculture.php",
    "CitiNewsroom": "https://citinewsroom.com/feed/",
    "FAO News (All topics)": "https://www.fao.org/news/rss/en",
    "FAO GIEWS": "https://www.fao.org/giews/rss/en/",
    "FreshPlaza Africa": "https://www.freshplaza.com/africa/rss.xml",
    "African Arguments": "https://africanarguments.org/feed/",
    "How We Made It In Africa": "https://www.howwemadeitinafrica.com/feed/",
    "Bizcommunity (Africa • Agri+Logistics)": "https://www.bizcommunity.com/GenerateRss.aspx?i=63,76&c=81",
}

# --- Impact Rules ---
IMPACT_RULES = {
    "Supply Risk": [
        r"\bexport (?:ban|restriction|control)\b",
        r"\b(?:import|trade) (?:ban|restriction|control)\b",
        r"\bembargo\b",
        r"\b(?:drought|flood|rainfall|weather)\b",
        r"\b(?:pest|disease|infestation)\b",
        r"\b(?:shortage|scarcity)\b",
        r"\b(?:strike|protest|unrest)\b",
        r"\bport (?:closure|congestion|delay)\b",
        r"\bharvest (?:delay|loss|damage)\b",
        r"\bproduction (?:issue|problem|concern)\b",
    ],
    "Price Upside": [
        r"\bstrong (?:demand|buying|interest)\b",
        r"\b(?:surge|spike|jump|rise|increase)\b",
        r"\b(?:grant|stimulus|support)\b",
        r"\b(?:incentive|subsidy|funding)\b",
        r"\bvalue[- ](?:addition|chain|processing)\b",
        r"\bhigh(?:er)? (?:price|demand|consumption)\b",
        r"\bmarket (?:rally|strength|upturn)\b",
        r"\bsupply (?:squeeze|shortage|tightness)\b",
        r"\bquality premium\b",
    ],
    "Price Downside": [
        r"\b(?:oversupply|surplus|glut)\b",
        r"\b(?:decline|fall|drop|decrease|slump)\b",
        r"\b(?:weak|soft|bearish) (?:price|market|demand)\b",
        r"\bcut (?:price|rate|cost)\b",
        r"\blow(?:er)? (?:price|demand|consumption)\b",
        r"\bmarket (?:weakness|downturn)\b",
        r"\bcompetitive pressure\b",
    ],
    "FX & Policy": [
        r"\b(?:depreciation|devaluation)\b",
        r"\bweak(?:ening)? (?:currency|exchange)\b",
        r"\b(?:fx|forex|dollar|euro|yuan)\b",
        r"\b(?:monetary|fiscal|trade) policy\b",
        r"\b(?:interest|exchange) rate\b",
        r"\b(?:tariff|duty|levy|tax)\b",
        r"\bregulatory (?:change|update|requirement)\b",
        r"\bpolicy (?:change|update|reform)\b",
    ],
    "Logistics & Trade": [
        r"\b(?:freight|shipping|transport)\b",
        r"\b(?:port|container|vessel|cargo)\b",
        r"\b(?:congestion|delay|bottleneck)\b",
        r"\b(?:reroute|divert|alternative route)\b",
        r"\b(?:cost|rate) (?:increase|surge|rise)\b",
        r"\btrade (?:flow|route|pattern)\b",
        r"\b(?:export|import) (?:volume|data|figure)\b",
    ],
    "Market Structure": [
        r"\b(?:merger|acquisition|takeover)\b",
        r"\b(?:investment|expansion|capacity)\b",
        r"\b(?:processing|factory|facility)\b",
        r"\b(?:certification|standard|quality)\b",
        r"\b(?:cooperative|association|group)\b",
        r"\b(?:contract|agreement|deal)\b",
        r"\b(?:partnership|collaboration)\b",
        r"\bmarket (?:structure|reform|development)\b",
    ],
    "Tech & Innovation": [
        r"\b(?:technology|innovation|digital)\b",
        r"\b(?:blockchain|traceability|tracking)\b",
        r"\b(?:sustainability|sustainable)\b",
        r"\b(?:efficiency|optimization)\b",
        r"\b(?:automation|mechanization)\b",
        r"\b(?:research|development|r&d)\b",
        r"\b(?:startup|fintech|agtech)\b",
    ],
}

# --- UI / CSS ---
CARD_CSS = """
<style>
.hero {
    position: relative;
    border-radius: 16px;
    padding: 28px 28px;
    background: linear-gradient(135deg, #0ea5e9, #7c3aed 60%);
    color: white;
    box-shadow: 0 14px 40px rgba(0,0,0,0.18);
}
.hero h1 { margin: 0 0 6px 0; font-size: 28px; font-weight: 800; }
.hero p { margin: 0; opacity: .95; }

.pill {
    display: inline-flex; align-items: center; gap: 8px;
    padding: 6px 12px; border-radius: 999px;
    background: rgba(255,255,255,0.15); color:#fff; font-weight:600; font-size: 13px;
}

.card {
    background: #ffffff;
    border: 1px solid rgba(0,0,0,.06);
    border-radius: 14px; overflow: hidden;
    transition: transform .15s ease, box-shadow .15s ease;
    height: 100%; /* NEW: Ensures cards in a row are same height */
    display: flex;
    flex-direction: column;
}
.card:hover { transform: translateY(-3px); box-shadow: 0 10px 24px rgba(0,0,0,0.08); }
.thumb { width: 100%; height: 180px; object-fit: cover; background:#f6f7f9; }
.card-body { padding: 14px; flex-grow: 1; display: flex; flex-direction: column; }
.card .title { color: #111827 !important; font-weight: 800; font-size: 18px; margin: 6px 0 8px 0; line-height: 1.25; }
.card .meta { color: #6b7280 !important; font-size: 12px; display:flex; gap:10px; flex-wrap:wrap; margin-bottom:8px; }
.card .summary { color:#374151 !important; font-size: 13px; line-height:1.55; margin-top: 6px; flex-grow: 1; }
.badges { display:flex; flex-wrap:wrap; gap:6px; margin:8px 0; }
.badge { font-size: 11px; font-weight:700; padding:4px 8px; border-radius:999px; background:#eef2ff; color:#3730a3; border:1px solid #c7d2fe; }
.link { text-decoration: none; font-weight:700; color:#2563eb !important; margin-top: 10px; }

/* NEW: Tweak tabs for a cleaner look */
[data-baseweb="tab-list"] {
    gap: 8px;
}
[data-baseweb="tab"] {
    background-color: #f6f7f9;
    border-radius: 8px 8px 0 0;
    padding-top: 10px;
    padding-bottom: 10px;
}
[data-baseweb="tab"]:hover {
    background-color: #f0f2f6;
}
[data-baseweb="tab"][aria-selected="true"] {
    background-color: #ffffff;
    border-bottom: 2px solid #7c3aed;
    font-weight: 600;
}
</style>
"""

# ==============================================================================
# 3. LOGGING & OPTIONAL LIBS
# ==============================================================================

st.set_option('client.showErrorDetails', False)

logger = logging.getLogger("oneafrica.pulse")
if not logger.handlers:
    handler = logging.StreamHandler()
    formatter = logging.Formatter("%(asctime)s - %(levelname)s - %(message)s")
    handler.setFormatter(formatter)
    logger.addHandler(handler)
logger.setLevel(logging.INFO)

if not HAS_SK:
    logger.info("sklearn not available; falling back to keyword hit scoring.")
if not HAS_VADER:
    logger.info("VADER not available; local sentiment analysis disabled.")
if not HAS_DOCX:
    logger.info("python-docx not available; .docx export will be disabled.")

# ==============================================================================
# 4. CORE HELPERS (SECRETS, HTTP, STATE, ERRORS)
# ==============================================================================

# --- Secrets ---
def _get_secret_safely(name: str) -> str:
    val = os.environ.get(name, "")
    if val:
        return str(val).strip().strip('"').strip("'")
    try:
        if hasattr(st, "secrets"):
            try:
                if len(st.secrets) > 0 and name in st.secrets:
                    return str(st.secrets.get(name, "")).strip().strip('"').strip("'")
            except Exception:
                pass
    except Exception:
        pass
    return ""

def get_newsdata_api_key() -> str:
    return _get_secret_safely("NEWSDATA_API_KEY")

def get_openai_api_key() -> str:
    return _get_secret_safely("OPENAI_API_KEY")

def get_twitter_bearer() -> str:
    return _get_secret_safely("TWITTER_BEARER_TOKEN") or _get_secret_safely("X_BEARER_TOKEN")

def get_xai_api_key() -> str:
    return _get_secret_safely("XAI_API_KEY")

# --- HTTP & Text Utils ---
@st.cache_resource(ttl=600)
def get_session() -> requests.Session:
    session = requests.Session()
    retries = Retry(total=3, backoff_factor=0.6, status_forcelist=[429, 500, 502, 503, 504])
    session.mount("http://", HTTPAdapter(max_retries=retries))
    session.mount("https://", HTTPAdapter(max_retries=retries))
    return session

def _normalize(text: str) -> str:
    text = html.unescape(text or "")
    text = re.sub(r"\s+", " ", text).strip()
    return text

def parse_date(date_str: str) -> Optional[dt.datetime]:
    try:
        fmts = [
            "%Y-%m-%dT%H:%M:%S%z", "%Y-%m-%dT%H:%M:%SZ",
            "%Y-%m-%d %H:%M:%S", "%a, %d %b %Y %H:%M:%S %z",
            "%a, %d %b %Y %H:%M:%S %Z", "%Y-%m-%d",
            "%d %b %Y", "%B %d, %Y",
        ]
        for fmt in fmts:
            try:
                return dt.datetime.strptime(date_str, fmt)
            except ValueError:
                continue
    except Exception as e:
        logger.info(f"parse_date: {e}")
    return None

def hash_key(*parts) -> str:
    return hashlib.md5(("||".join([p or "" for p in parts])).encode("utf-8")).hexdigest()

# --- Session State ---
def ss_get(key, default):
    if key not in st.session_state:
        st.session_state[key] = default
    return st.session_state[key]

def ss_set(key, value):
    st.session_state[key] = value

def init_chat_state():
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = [
            {"role": "system", "content": (
                "You are a crisp market-intelligence assistant for West African tree crops "
                "(cashew, shea, cocoa, palm kernel). Be concise, cite assumptions, and suggest "
                "actionable next steps. If asked to summarize a table, write bullet points."
            )}
        ]

# --- Error Handling ---
SOFT_ERRORS: List[str] = []
def soft_fail(msg: str, detail: Optional[str] = None):
    if msg:
        SOFT_ERRORS.append(msg)
    if detail:
        logger.warning(detail)

def friendly_error_summary():
    if not SOFT_ERRORS:
        return
    counts: Dict[str,int] = {}
    for m in SOFT_ERRORS:
        counts[m] = counts.get(m, 0) + 1
    bullets = "".join([f"- {msg} _(x{n})_\n" for msg, n in counts.items()])
    st.info(f"""
    **Heads up:** Some sources were temporarily skipped or partially loaded.
    This doesn’t affect your ability to scan and summarize current items.

    {bullets}
    """)
    SOFT_ERRORS.clear() # Clear after displaying

# ==============================================================================
# 5. AI & LLM FUNCTIONS (OPENAI, GROK)
# ==============================================================================

# --- OpenAI Helpers ---
def have_openai():
    """Helper to check if OpenAI lib is loaded and key is present."""
    return OPENAI_OK and bool(get_openai_api_key())

@st.cache_resource(ttl=3600)
def get_openai_client():
    """Robust client factory that respects OPENAI_BASE_URL if set."""
    try:
        api_key = get_openai_api_key()
        if not api_key:
            return None
        base_url = os.environ.get("OPENAI_BASE_URL", "").strip()
        if base_url:
            return OpenAI(api_key=api_key, base_url=base_url)
        return OpenAI(api_key=api_key)  # official endpoint
    except Exception as e:
        logger.warning(f"OpenAI client init failed: {e}")
        return None

# --- Grok/xAI Helpers ---
def have_grok():
    """Helper to check if Grok lib is loaded and key is present."""
    return OPENAI_OK and bool(get_xai_api_key())

@st.cache_resource(ttl=3600)
def get_grok_client():
    """Robust client factory for Grok."""
    try:
        api_key = get_xai_api_key()
        if not api_key:
            return None
        return OpenAI(api_key=api_key, base_url="https://api.x.ai/v1")
    except Exception as e:
        logger.warning(f"Grok client init failed: {e}")
        return None

# --- LLM Summary ---
@st.cache_data(ttl=60*60, show_spinner=False) # Cache for 1 hour
def llm_abstractive_summary(text: str, keywords: List[str]) -> str:
    """
    Generates a concise, abstractive summary using an LLM, focused on
    the provided keywords.
    """
    if not text or not have_openai():
        return ""
    client = get_openai_client()
    if client is None:
        return ""

    kw_str = ", ".join(keywords)
    prompt = f"""
    You are a market analyst. Summarize the following article text in 2-3 concise sentences.
    Focus *only* on facts relevant to these keywords: {kw_str}
    If the text has no relevance to the keywords, return an empty string.

    Article (may be partial):
    {text[:7000]}
    """
    try:
        resp = client.chat.completions.create(
            model="gpt-4o-mini", # Use a fast, modern model
            temperature=0.1,
            messages=[
                {"role": "system", "content": "Be concise, factual, and strictly relevant to the provided keywords."},
                {"role": "user", "content": prompt},
            ],
        )
        summary = (resp.choices[0].message.content or "").strip()
        return summary
    except Exception as e:
        logger.warning(f"llm_abstractive_summary failed: {e}")
        soft_fail("One LLM summary failed, used fallback.", f"llm_summary EXC {e}")
        return ""

# --- LLM Article Analysis ---
@st.cache_data(ttl=30*60, show_spinner=False)
def _llm_analyze_article_cached(model: str, title: str, body: str, tags: List[str]) -> str:
    client = get_openai_client()
    if client is None:
        return ""
    prompt = f"""
You are a market-intelligence analyst focused on tree crop commodities in the world
(cashew, shea, cocoa, palm kernel), logistics, and FX.

Analyze the ARTICLE and produce a concise, executive-ready brief. Use short, punchy bullets
where appropriate and provide concrete, actionable guidance. Avoid fluff.

ARTICLE TITLE:
{title[:400]}

ARTICLE BODY (may be partial):
{body[:7000]}

HEURISTIC TAGS PROVIDED BY UI (may be incomplete):
{", ".join(tags) if tags else "General"}

Return your analysis in EXACTLY these sections with clear headings:
1) WHAT THE ARTICLE MEANS — 2–3 sentence synthesis
2) KEY INSIGHTS — 3–6 bullets with the most important takeaways
3) MARKET IMPACT — specific effects on supply/demand, prices, logistics, FX; note direction & magnitude if possible
4) BUSINESS OPPORTUNITIES — 3–6 concrete moves we could make now (be specific)
5) RISK FACTORS — 3–5 concise bullets (operational, financial/FX, regulatory)
6) ACTIONABLE RECOMMENDATIONS — 3–5 steps with owners or thresholds where relevant
7) TIME HORIZON — near-term (0–3m) / medium (3–12m) / long (12m+)
8) CONFIDENCE — High/Medium/Low and why

Constraints:
- Keep it pragmatic and West-Africa oriented.
- If information is uncertain, say so explicitly and suggest a verification step.
"""
    resp = client.chat.completions.create(
        model=model, temperature=0.3,
        messages=[
            {"role": "system", "content": "Be precise, actionable, and bias towards decisions and thresholds."},
            {"role": "user", "content": prompt},
        ],
    )
    return (resp.choices[0].message.content or "").strip()

def analyze_with_llm(title: str, body: str, tags: List[str]) -> str:
    if not have_openai():
        return ""
    model_candidates = ["gpt-4o-mini","gpt-4o","gpt-3.5-turbo-0125"]
    for m in model_candidates:
        try:
            out = _llm_analyze_article_cached(m, title, body, tags)
            if out:
                return out
        except Exception as e:
            logger.warning(f"LLM analyze failed on {m}: {e}")
            continue
    return ""

# --- LLM Chat Assistant ---
def generate_assistant_reply(messages, temperature: float = 0.4):
    if not have_openai():
        return None, False
    client = get_openai_client()
    if client is None:
        return None, False

    model_candidates = [
        "gpt-4o-mini",
        "gpt-4o",
        "gpt-3.5-turbo-0125",
    ]
    last_err = None
    
    # Try streaming first
    for model in model_candidates:
        try:
            stream = client.chat.completions.create(
                model=model, messages=messages, stream=True, temperature=temperature
            )
            chunks = []
            with st.chat_message("assistant"):
                placeholder = st.empty()
                buf = ""
                for chunk in stream:
                    delta = chunk.choices[0].delta.content or ""
                    if delta:
                        buf += delta
                        placeholder.markdown(buf + "▌") # Add cursor
                placeholder.markdown(buf) # Final message
                chunks.append(buf)
            reply = "".join(chunks).strip()
            if reply:
                return reply, True
        except Exception as e:
            logger.warning(f"OpenAI streaming failed on {model}: {e}")
            last_err = e
            continue # Try next model
            
    # Fallback to non-streaming
    for model in model_candidates:
        try:
            comp = client.chat.completions.create(
                model=model, messages=messages, temperature=temperature
            )
            reply = (comp.choices[0].message.content or "").strip()
            if reply:
                return reply, False # Not streamed
        except Exception as e2:
            logger.warning(f"OpenAI non-streaming failed on {model}: {e2}")
            last_err = e2
            continue
            
    soft_fail("Assistant is temporarily unavailable.", f"OpenAI failures: {last_err}")
    return None, False

# ==============================================================================
# 6. DATA FETCHING (RSS, NEWSDATA, TWITTER)
# ==============================================================================

# --- Web Page Fetching ---
@st.cache_data(ttl=60*30, show_spinner=False)
def fetch_page(url: str, timeout: int = 12) -> str:
    try:
        headers = {
            "User-Agent": "Mozilla/5.0 (compatible; OneAfricaPulse/1.0)",
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
            "Accept-Language": "en-US,en;q=0.5",
        }
        r = get_session().get(url, headers=headers, timeout=timeout)
        if r.status_code != 200:
            soft_fail("Skipped a page that didn’t load cleanly.", f"fetch_page {url} -> {r.status_code}")
            return ""
        return r.text
    except Exception as e:
        soft_fail("Skipped one page due to connectivity.", f"fetch_page EXC {url}: {e}")
        return ""

# --- RSS/Atom Fetching ---
ATOM_NS = "{http://www.w3.org/2005/Atom}"

def _text(elem: Optional[ET.Element]) -> str:
    return _normalize(elem.text if elem is not None and elem.text else "")

def _find(elem: ET.Element, tag: str) -> Optional[ET.Element]:
    e = elem.find(tag)
    if e is not None:
        return e
    if not tag.startswith("{"):
        e = elem.find(ATOM_NS + tag)
    return e

def _findall(elem: ET.Element, tag: str) -> List[ET.Element]:
    return list(elem.findall(tag)) + list(elem.findall(ATOM_NS + tag))

@st.cache_data(ttl=60*10, show_spinner=False)
def fetch_feed_raw(url: str, timeout: int = 20) -> bytes:
    try:
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) Chrome/124.0 OneAfricaPulse/1.0",
            "Accept": "application/rss+xml, application/atom+xml, application/xml, text/xml, */*",
        }
        r = get_session().get(url, headers=headers, timeout=timeout, allow_redirects=True)
        if r.status_code != 200:
            soft_fail("Skipped a source that returned a non-200 response.", f"fetch_feed_raw {url} -> {r.status_code}")
        return r.content if r.status_code == 200 else (r.content or b"")
    except Exception as e:
        soft_fail("Temporarily skipped one source due to connectivity.", f"fetch_feed_raw EXC {e}")
        return b""

def parse_feed_xml(content: bytes, base_url: str) -> List[Dict[str, str]]:
    items: List[Dict[str, str]] = []
    if not content:
        return items
    try:
        sample = content.lstrip()[:16]
        if not sample.startswith(b"<") and not sample.startswith(b"\xef\xbb\xbf<"):
            soft_fail("Skipped one feed that returned non-XML.", f"parse_feed_xml non-xml from {base_url}")
            return items
    except Exception:
        pass
    try:
        root = ET.fromstring(content)
        channel = root.find("channel")
        if channel is not None:  # RSS
            for it in channel.findall("item"):
                title = _text(_find(it, "title")) or "(untitled)"
                link = _text(_find(it, "link"))
                if not link:
                    guid = _find(it, "guid")
                    link = _text(guid)
                if link and link.startswith("/"):
                    link = urljoin(base_url, link)
                summary = _text(_find(it, "description"))
                pub = _text(_find(it, "pubDate"))
                if title or link:
                    items.append({"title": title, "link": link, "summary": summary, "published_raw": pub})
            return items

        # Atom
        for entry in _findall(root, "entry"):
            title = _text(_find(entry, "title")) or "(untitled)"
            link_el = _find(entry, "link")
            link = ""
            if link_el is not None:
                link = link_el.attrib.get("href", "") or _text(link_el)
            if link and link.startswith("/"):
                link = urljoin(base_url, link)
            summary = _text(_find(entry, "summary")) or _text(_find(entry, "content"))
            pub = _text(_find(entry, "updated")) or _text(_find(entry, "published"))
            if title or link:
                items.append({"title": title, "link": link, "summary": summary, "published_raw": pub})
        return items
    except Exception as e:
        soft_fail("Skipped one feed that had invalid XML.", f"parse_feed_xml EXC {e}")
        return items

def validate_feed(url: str, ignore_recency_check: bool = False) -> Tuple[bool, str]:
    try:
        content = fetch_feed_raw(url)
        items = parse_feed_xml(content, base_url=url)
        if not items:
            return False, "No entries found"
        if ignore_recency_check:
            return True, "OK"
        now = dt.datetime.now(dt.timezone.utc)
        for it in items[:10]:
            d = parse_date(it.get("published_raw", "") or "")
            if d:
                d = d.astimezone(dt.timezone.utc) if d.tzinfo else d.replace(tzinfo=dt.timezone.utc)
                if (now - d).days <= 60:
                    return True, "OK"
        return False, "No recent entries (≤60 days)"
    except Exception as ex:
        soft_fail("Skipped one feed due to a validation issue.", f"validate_feed EXC {url}: {ex}")
        return False, "Validation error"

def fetch_from_feed(url: str, start_date: dt.datetime, end_date: dt.datetime,
                    force_fetch: bool, ignore_recency: bool) -> List[Dict[str, Any]]:
    ok, status = validate_feed(url, ignore_recency_check=ignore_recency)
    if not ok and not force_fetch:
        soft_fail(f"Skipped {urlparse(url).netloc} (feed not recent/valid).", f"validate -> {status}")
        return []
    raw = fetch_feed_raw(url)
    raw_items = parse_feed_xml(raw, base_url=url)

    items: List[Dict[str, Any]] = []
    for e in raw_items:
        title = _normalize(e.get("title", ""))
        link = e.get("link", "")
        summary = _normalize(e.get("summary", ""))
        published = None
        if e.get("published_raw"):
            published = parse_date(e["published_raw"])
        if published:
            published = published.astimezone(dt.timezone.utc) if published.tzinfo else published.replace(tzinfo=dt.timezone.utc)
            if not (start_date <= published <= end_date):
                continue
            published_str = published.strftime("%Y-%m-%d %H:%M UTC")
        else:
            # If no date, we can't filter, so we assume it's recent enough for the feed.
            # This is a tradeoff for feeds without pubDates.
            published_str = "Date unknown"

        items.append({
            "source": urlparse(url).netloc,
            "title": title,
            "link": link,
            "published": published_str,
            "summary": summary,
        })
    return items

# --- Newsdata.io Fetching ---
NEWSDATA_BASE = "https://newsdata.io/api/1/latest"

@st.cache_data(ttl=60*10, show_spinner=False)
def fetch_from_newsdata_cached(redacted_params: Dict[str, Any], max_pages: int) -> List[Dict[str, Any]]:
    # This function's sole purpose is to be a cache placeholder.
    # The actual work is done by the runtime function.
    return []

def fetch_from_newsdata_runtime(api_key: str, base_params: Dict[str, Any], max_pages: int) -> List[Dict[str, Any]]:
    session = get_session()
    items: List[Dict[str, Any]] = []
    pages = 0
    next_page = None

    params = dict(base_params)
    params["apikey"] = api_key

    while pages < max_pages:
        try:
            q = dict(params)
            if next_page:
                q["page"] = next_page
            r = session.get(NEWSDATA_BASE, params=q, timeout=20)
            if r.status_code != 200:
                soft_fail("One API page was skipped (non-200).", f"newsdata {r.status_code} {r.text[:200]}")
                break
            data = r.json()
            results = data.get("results") or data.get("articles") or []
            for a in results:
                items.append(a)
            next_page = data.get("nextPage") or data.get("next_page")
            pages += 1
            if not next_page:
                break
            time.sleep(0.5) # Rate limit
        except Exception as e:
            soft_fail("Temporarily skipped an API page due to connectivity.", f"newsdata EXC {e}")
            break
    return items

def fetch_from_newsdata(
    api_key: str,
    query: str,
    start_date: dt.datetime,
    end_date: dt.datetime,
    language: Optional[str] = None,
    country: Optional[str] = None,
    category: Optional[str] = None,
    max_pages: int = 2,
) -> List[Dict[str, Any]]:
    if not api_key:
        return []
    redacted = {"q": query or ""}
    if language: redacted["language"] = language
    if country: redacted["country"] = country
    if category: redacted["category"] = category
    
    # Use date range for query, as API is 'latest'
    redacted["from_date"] = start_date.strftime("%Y-%m-%d")
    redacted["to_date"] = end_date.strftime("%Y-%m-%d")

    _ = fetch_from_newsdata_cached(redacted, max_pages=max_pages)
    items_raw = fetch_from_newsdata_runtime(api_key=api_key, base_params=redacted, max_pages=max_pages)

    items: List[Dict[str, Any]] = []
    for a in items_raw:
        try:
            title = _normalize(a.get("title", ""))
            link = a.get("link") or a.get("url") or ""
            source = a.get("source_id") or a.get("source") or "newsdata.io"
            pub = a.get("pubDate") or a.get("published_at") or ""
            desc = _normalize(a.get("description", "")) or _normalize(a.get("content", ""))

            ok_date = True
            published_str = "Date unknown"
            if pub:
                d = parse_date(pub)
                if d:
                    d = d.astimezone(dt.timezone.utc) if d.tzinfo else d.replace(tzinfo=dt.timezone.utc)
                    # We already filtered by API, but double-check
                    ok_date = start_date <= d <= end_date
                    published_str = d.strftime("%Y-%m-%d %H:%M UTC")
            if not ok_date:
                continue
                
            items.append({
                "source": f"{source} (newsdata.io)",
                "title": title or "(untitled)",
                "link": link,
                "published": published_str,
                "summary": desc,
            })
        except Exception as e:
            soft_fail("Skipped one API item due to missing fields.", f"newsdata item EXC {e}")
            continue
    return items

# --- Twitter/X Fetching (Grok) ---
def fetch_tweets_via_grok(query: str, lang: str, hours: int, max_results: int = 300) -> List[Dict[str, Any]]:
    client = get_grok_client()
    if not client:
        soft_fail("XAI_API_KEY missing for Grok integration.", "Set XAI_API_KEY in .env or secrets.")
        return []

    prompt = f"""
You are a helpful assistant that can access real-time X (Twitter) posts.
Fetch up to {max_results} recent posts matching this query: {query} lang:{lang}
Look back no more than {hours} hours.
Return ONLY a JSON array of objects with these fields for each post:
- id: string
- created_at: ISO format (YYYY-MM-DDTHH:MM:SSZ)
- text: full text
- lang: language code
- retweets: integer
- likes: integer
- username: screen name
- url: full URL like "https://x.com/username/status/id"
No explanations or other text.
"""
    try:
        response = client.chat.completions.create(
            model="grok-beta",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.0,
            max_tokens=4096,
        )
        raw_content = response.choices[0].message.content.strip()
        # Find the JSON blob
        json_match = re.search(r'\[.*\]', raw_content, re.DOTALL)
        if not json_match:
            soft_fail("Grok tweet fetch failed (no JSON).", f"Grok response: {raw_content[:200]}")
            return []
        
        tweets = json.loads(json_match.group(0))
        return tweets
    except Exception as e:
        soft_fail("Grok tweet fetch failed.", f"EXC: {e}")
        return []

# --- Twitter/X Fetching (v2 API) ---
@st.cache_data(ttl=60*10, show_spinner=False)
def fetch_tweets_via_twitter_v2(bearer_token: str, query: str, lang: str, start_time_iso: str, max_results: int) -> List[Dict[str, Any]]:
    if not bearer_token:
        soft_fail("TWITTER_BEARER_TOKEN missing.", "Set TWITTER_BEARER_TOKEN in .env or secrets.")
        return []

    session = get_session()
    url = "https://api.twitter.com/2/tweets/search/recent"
    
    full_query = query
    if lang:
        full_query += f" lang:{lang}"
    
    params = {
        "query": full_query,
        "max_results": max(10, min(100, max_results)), # Twitter max_results is 10-100
        "start_time": start_time_iso,
        "tweet.fields": "created_at,lang,public_metrics",
        "expansions": "author_id",
        "user.fields": "username",
    }
    headers = {
        "Authorization": f"Bearer {bearer_token}",
        "User-Agent": "OneAfricaPulse/1.0",
    }
    
    try:
        r = session.get(url, params=params, headers=headers, timeout=20)
        if r.status_code != 200:
            soft_fail("Twitter v2 API fetch failed (non-200).", f"Twitter v2 API {r.status_code}: {r.text[:200]}")
            return []
        
        data = r.json()
        tweets = data.get("data", [])
        includes = data.get("includes", {})
        users = {u["id"]: u["username"] for u in includes.get("users", [])}
        
        results = []
        for t in tweets:
            author_id = t.get("author_id")
            username = users.get(author_id, "unknown")
            metrics = t.get("public_metrics", {})
            results.append({
                "id": t.get("id"),
                "created_at": t.get("created_at"),
                "text": t.get("text"),
                "lang": t.get("lang"),
                "retweets": metrics.get("retweet_count", 0),
                "likes": metrics.get("like_count", 0),
                "username": username,
                "url": f"https://x.com/{username}/status/{t.get('id')}",
            })
        return results
    except Exception as e:
        soft_fail("Twitter v2 API request failed.", f"EXC: {e}")
        return []

# ==============================================================================
# 7. DATA PROCESSING (ENRICHMENT, SENTIMENT, ORCHESTRATION)
# ==============================================================================

# --- Article Content Extraction ---
def get_og_image(soup: BeautifulSoup, base_url: str) -> Optional[str]:
    try:
        candidates = [
            ("meta", {"property": "og:image"}),
            ("meta", {"name": "twitter:image"}),
            ("meta", {"property": "twitter:image"}),
            ("link", {"rel": "image_src"}),
        ]
        for tag, attrs in candidates:
            el = soup.find(tag, attrs=attrs)
            if el:
                src = el.get("content") or el.get("href")
                if src:
                    if src.startswith("//"):
                        return "https:" + src
                    if src.startswith("/"):
                        return urljoin(base_url, src)
                    return src
    except Exception as e:
        logger.info(f"get_og_image: {e}")
    return None

def get_favicon_url(domain_url: str) -> str:
    try:
        parsed = urlparse(domain_url)
        return f"{parsed.scheme}://{parsed.netloc}/favicon.ico"
    except Exception:
        return ""

@st.cache_data(ttl=60*30, show_spinner=False)
def fetch_article_text_and_image(url: str) -> Tuple[str, str]:
    if not url:
        return "", FALLBACK_IMG
    html_text = fetch_page(url)
    if not html_text:
        return "", FALLBACK_IMG
    try:
        soup = BeautifulSoup(html_text, "html.parser")
        for tag in soup(["script", "style", "noscript", "nav", "footer", "iframe", "form", "header"]):
            tag.decompose()

        candidates = []
        selectors = [
            "article", "[role='main']", ".article-body", ".story-body",
            ".post-content", "main", ".content", ".entry-content",
            "#article-body", ".article-content", ".story-content",
            ".news-content", ".page-content", "body",
        ]
        for sel in selectors:
            for node in soup.select(sel):
                text = node.get_text(separator=" ", strip=True)
                if len(text) > 100:
                    candidates.append(text)
        text = max(candidates, key=len) if candidates else soup.get_text(separator=" ", strip=True)
        text = _normalize(text)
        if len(text) < 50:
            text = ""

        img = get_og_image(soup, url) or get_favicon_url(url) or FALLBACK_IMG
        return text, img
    except Exception as e:
        soft_fail("Used a fallback image for one article.", f"fetch_article_text_and_image EXC {url}: {e}")
        return "", FALLBACK_IMG

# --- Relevance & Summary (Local) ---
def keyword_relevance(text: str, keywords: List[str]) -> float:
    if not text:
        return 0.0
    if HAS_SK:
        try:
            vec = TfidfVectorizer(stop_words="english", max_features=5000)
            X = vec.fit_transform([text, " ".join(keywords)])
            sim = cosine_similarity(X[0:1], X[1:2])[0][0]
            return float(sim)
        except Exception as e:
            logger.info(f"tfidf relevance fallback: {e}")
    # Fallback
    tokens = re.findall(r"[a-zA-Z']{3,}", text.lower())
    kwset = {k.lower() for k in keywords}
    hits = sum(1 for t in tokens if t in kwset)
    return hits / max(1, len(tokens))

def simple_extractive_summary(text: str, n_sentences: int = 3, keywords: Optional[List[str]] = None) -> str:
    if not text:
        return ""
    sents = re.split(r"(?<=[\.\?\!])\s+", text)
    sents = [s for s in sents if 30 <= len(s) <= 400][:60]
    if len(sents) <= n_sentences:
        return " ".join(sents)
    if HAS_SK:
        try:
            vec = TfidfVectorizer(stop_words="english", max_features=8000)
            X = vec.fit_transform(sents)  # sparse
            centroid = np.asarray(X.mean(axis=0)).ravel()
            sims = cosine_similarity(X, centroid.reshape(1, -1)).ravel()
            if keywords:
                kw = [k.lower() for k in keywords]
                boost = np.array([sum(1 for w in re.findall(r"[a-z']+", s.lower()) if w in kw) for s in sents], dtype=float)
                sims = sims + 0.05 * boost
            idx = sims.argsort()[-n_sentences:][::-1]
            return " ".join([sents[i] for i in idx])
        except Exception as e:
            logger.info(f"summary fallback: {e}")
    return " ".join(sents[:n_sentences])

# --- Impact Classification ---
def classify_impact(text: str) -> List[str]:
    tags = []
    lower = text.lower()
    for label, patterns in IMPACT_RULES.items():
        try:
            if any(re.search(p, lower) for p in patterns):
                tags.append(label)
        except Exception:
            continue
    return list(dict.fromkeys(tags)) or ["General"]

# --- Sentiment Analysis (Grok) ---
def analyze_tweet_sentiment_grok(tweets: List[Dict[str, Any]]) -> pd.DataFrame:
    client = get_grok_client()
    if not client:
        return pd.DataFrame()

    texts = [t['text'] for t in tweets]
    batch_size = 50  # to avoid token limits
    compounds = []
    for i in range(0, len(texts), batch_size):
        batch = texts[i:i+batch_size]
        prompt = """
You are a sentiment analyst. For each text, return a compound score from -1 (negative) to +1 (positive).
Return ONLY a JSON array of floats, no explanations.
Texts:
""" + "\n---\n".join(batch)
        try:
            response = client.chat.completions.create(
                model="grok-beta",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.0,
                max_tokens=1000,
            )
            raw = response.choices[0].message.content.strip()
            json_match = re.search(r'\[.*\]', raw, re.DOTALL)
            if not json_match:
                raise Exception("No JSON array found in Grok response")
            scores = json.loads(json_match.group(0))
            if len(scores) != len(batch):
                raise Exception(f"Score count ({len(scores)}) != batch count ({len(batch)})")
            compounds.extend(scores)
        except Exception as e:
            soft_fail("Grok sentiment analysis failed.", f"EXC: {e}")
            compounds.extend([0.0] * len(batch))

    rows = []
    for t, sc in zip(tweets, compounds):
        label = "Neutral"
        if sc >= 0.05:
            label = "Positive"
        elif sc <= -0.05:
            label = "Negative"
        rows.append({
            "created_at": t.get("created_at"),
            "text": t.get("text", ""),
            "compound": sc,
            "label": label,
            "retweets": t.get("retweets", 0),
            "likes": t.get("likes", 0),
            "username": t.get("username", ""),
            "url": t.get("url", ""),
        })
    df = pd.DataFrame(rows)
    if not df.empty:
        try:
            df["created_at"] = pd.to_datetime(df["created_at"])
            df = df.sort_values("created_at", ascending=False).reset_index(drop=True)
        except Exception:
            pass
    return df

# --- Sentiment Analysis (VADER) ---
@st.cache_resource
def get_sentiment_analyzer():
    if not HAS_VADER:
        return None
    try:
        return SentimentIntensityAnalyzer()
    except Exception as e:
        soft_fail("VADER analyzer unavailable.", f"vader EXC {e}")
        return None

def analyze_tweet_sentiment(tweets: List[Dict[str, Any]]) -> pd.DataFrame:
    if not tweets:
        return pd.DataFrame(columns=["created_at","text","compound","label","retweets","likes","username","url"])
    sia = get_sentiment_analyzer()
    if not sia:
        soft_fail("VADER sentiment analysis disabled.", "VADER lib missing or failed to init.")
        return pd.DataFrame()
        
    rows = []
    for t in tweets:
        text = _normalize(t.get("text",""))
        if not text:
            continue
        sc = sia.polarity_scores(text).get("compound", 0.0)
        label = "Neutral"
        if sc >= 0.05:
            label = "Positive"
        elif sc <= -0.05:
            label = "Negative"
        rows.append({
            "created_at": t.get("created_at"),
            "text": text,
            "compound": sc,
            "label": label,
            "retweets": t.get("retweets", 0),
            "likes": t.get("likes", 0),
            "username": t.get("username", ""),
            "url": t.get("url", ""),
        })
    df = pd.DataFrame(rows)
    if not df.empty:
        try:
            df["created_at"] = pd.to_datetime(df["created_at"])
            df = df.sort_values("created_at", ascending=False).reset_index(drop=True)
        except Exception:
            pass
    return df

def summarize_sentiment(df: pd.DataFrame) -> Dict[str, Any]:
    if df is None or df.empty:
        return {"n": 0, "mean_compound": 0.0, "share_pos": 0.0, "share_neu": 1.0, "share_neg": 0.0}
    n = len(df)
    mean_c = float(df["compound"].mean())
    share_pos = float((df["label"] == "Positive").mean())
    share_neu = float((df["label"] == "Neutral").mean())
    share_neg = float((df["label"] == "Negative").mean())
    return {"n": n, "mean_compound": mean_c, "share_pos": share_pos, "share_neu": share_neu, "share_neg": share_neg}

# --- Core Enrichment & Orchestration ---
def enrich(entry: Dict[str, Any], keywords: List[str], min_relevance: float, n_sent: int, summary_method: str) -> Optional[Dict[str, Any]]:
    try:
        article_text, image_url = fetch_article_text_and_image(entry.get("link",""))
        base = entry.get("summary") or ""
        body = article_text if len(article_text) > len(base) else base

        rel = keyword_relevance(" ".join([entry.get("title",""), body]), keywords)
        if rel < min_relevance:
            return None
        
        summary = ""
        if summary_method == "Abstractive (LLM)":
            if not have_openai():
                soft_fail("LLM Summary skipped (no API key).", "LLM Summary skipped")
                summary = simple_extractive_summary(body, n_sentences=n_sent, keywords=keywords)
            else:
                summary = llm_abstractive_summary(body, keywords)
                if not summary: # Fallback if LLM fails or returns empty
                    summary = simple_extractive_summary(body, n_sentences=n_sent, keywords=keywords)
        else: # Default to extractive
            summary = simple_extractive_summary(body, n_sentences=n_sent, keywords=keywords)

        impacts = classify_impact(" ".join([entry.get("title",""), body])) or ["General"]

        return {
            "source": entry.get("source",""),
            "title": entry.get("title","(untitled)"),
            "link": entry.get("link",""),
            "published": entry.get("published","Date unknown"),
            "relevance": float(rel),
            "impact": impacts,
            "auto_summary": summary,
            "image": image_url or FALLBACK_IMG,
        }
    except Exception as e:
        soft_fail("Skipped one article that couldn’t be processed.", f"enrich EXC {e}")
        return None

def process_rows(rows: List[Dict[str, Any]]) -> pd.DataFrame:
    if not rows:
        return pd.DataFrame(columns=["source","published","title","relevance","impact","auto_summary","link","image"])
    
    # Simple dedupe by key
    seen = set()
    cleaned = []
    for r in rows:
        key = hash_key(r.get("title",""), r.get("link",""))
        if key in seen:
            continue
        seen.add(key)
        cleaned.append(r)
        
    df = pd.DataFrame(cleaned)
    if df.empty:
        return pd.DataFrame(columns=["source","published","title","relevance","impact","auto_summary","link","image"])
    return df.sort_values("relevance", ascending=False).reset_index(drop=True)

def fetch_all(params: Dict[str, Any]) -> List[Dict[str, Any]]:
    rows: List[Dict[str, Any]] = []
    total_tasks = len(params["chosen_sources"]) + (1 if (params["use_newsdata"] and params["newsdata_key"]) else 0)
    total_tasks = max(total_tasks, 1)
    
    # Use a shared progress bar in the sidebar
    progress_bar = st.sidebar.progress(0.0)
    info_text = st.sidebar.empty()

    # 1) RSS/Atom
    for i, src in enumerate(params["chosen_sources"], start=1):
        info_text.info(f"Fetching RSS {i}/{total_tasks}: {urlparse(src).netloc}")
        try:
            raw_items = fetch_from_feed(src, params["start_date"], params["end_date"], params["force_fetch"], params["ignore_recency"])
        except Exception as e:
            soft_fail("Skipped a source due to a transient issue.", f"fetch_from_feed EXC {src}: {e}")
            raw_items = []
        if params["per_source_cap"] and raw_items:
            raw_items = raw_items[:params["per_source_cap"]]

        if raw_items:
            with ThreadPoolExecutor(max_workers=8) as ex:
                futures = [ex.submit(enrich,
                                     {**e, "source": urlparse(src).netloc},
                                     params["keywords"],
                                     params["min_relevance"],
                                     params["n_sent"],
                                     params["summary_method"]
                                    ) for e in raw_items]
                for fut in as_completed(futures):
                    try:
                        r = fut.result()
                        if r: rows.append(r)
                    except Exception as e:
                        soft_fail("One article was skipped during processing.", f"future enrich EXC {e}")
        progress_bar.progress(min(1.0, i / total_tasks))

    # 2) Newsdata.io
    if params["use_newsdata"] and params["newsdata_key"]:
        info_text.info(f"Fetching Newsdata.io {len(params['chosen_sources'])+1}/{total_tasks}")
        try:
            nd_items = fetch_from_newsdata(
                api_key=params["newsdata_key"],
                query=params["newsdata_query"],
                start_date=params["start_date"],
                end_date=params["end_date"],
                language=params["nd_language"] or None,
                country=params["nd_country"] or None,
                category=params["nd_category"] or None,
                max_pages=int(params["nd_pages"]),
            )
            if params["per_source_cap"] and nd_items:
                nd_items = nd_items[:params["per_source_cap"]]
            if nd_items:
                with ThreadPoolExecutor(max_workers=8) as ex:
                    futures = [ex.submit(enrich,
                                         it,
                                         params["keywords"],
                                         params["min_relevance"],
                                         params["n_sent"],
                                         params["summary_method"]
                                        ) for it in nd_items]
                    for fut in as_completed(futures):
                        try:
                            r = fut.result()
                            if r: rows.append(r)
                        except Exception as e:
                            soft_fail("One API article was skipped during processing.", f"future API enrich EXC {e}")
        except Exception as e:
            soft_fail("The API was briefly unavailable; results shown are from RSS.", f"fetch_from_newsdata EXC {e}")
        progress_bar.progress(1.0)

    info_text.empty()
    progress_bar.empty()

    # Optional cross-source de-duplication
    if params.get("dedupe", True) and rows:
        seen = set()
        deduped = []
        for r in rows:
            k = hash_key(r["title"], r["link"])
            if k in seen:
                continue
            seen.add(k)
            deduped.append(r)
        rows = deduped
    return rows

# ==============================================================================
# 8. EXPORTING (DOCX, MARKDOWN)
# ==============================================================================

# --- DOCX Builder ---
def _set_heading_style(p, size=16, bold=True):
    run = p.runs[0] if p.runs else p.add_run()
    run.font.size = Pt(size)
    run.font.bold = bold

def _add_kv_para(doc: Document, key: str, val: str):
    p = doc.add_paragraph()
    p.add_run(f"{key}: ").bold = True
    p.add_run(val or "-")

def build_results_docx(
    app_name: str,
    tagline: str,
    quote: str,
    results_df: pd.DataFrame,
    digest_md: str,
    params: Dict[str, Any],
    sent_summary: Dict[str, Any] | None = None
) -> bytes:
    if not HAS_DOCX:
        soft_fail("Cannot generate .docx, `python-docx` not installed.", "pip install python-docx")
        return b""
        
    doc = Document()

    # --- Cover / Title ---
    title = doc.add_paragraph(app_name)
    _set_heading_style(title, size=22, bold=True)
    sub = doc.add_paragraph(tagline)
    _set_heading_style(sub, size=12, bold=False)
    doc.add_paragraph(quote)

    ts = dt.datetime.now(dt.timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    _add_kv_para(doc, "Generated", ts)
    try:
        _add_kv_para(
            doc, "Date Window",
            f"{params.get('start_date').strftime('%Y-%m-%d')} → {params.get('end_date').strftime('%Y-%m-%d')}"
        )
    except Exception:
        pass
    _add_kv_para(doc, "Min Relevance", f"{params.get('min_relevance',0):.2f}")
    _add_kv_para(doc, "Keywords", ", ".join(params.get('keywords', [])) or "-")
    doc.add_paragraph().add_run(" ").add_break()

    # --- Results summary table ---
    doc.add_paragraph("Results Overview").runs[0].bold = True
    if results_df is None or results_df.empty:
        doc.add_paragraph("No relevant items found for the selected period.")
    else:
        show_cols = ["source", "published", "title", "relevance", "impact"]
        tmp = results_df[show_cols].copy()
        tmp["impact"] = tmp["impact"].apply(lambda x: ", ".join(x) if isinstance(x, list) else str(x))
        tmp["relevance"] = tmp["relevance"].apply(lambda r: f"{r:.0%}")

        table = doc.add_table(rows=1, cols=len(show_cols))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, c in enumerate(["Source", "Published", "Title", "Relevance", "Impact"]):
            hdr_cells[i].text = c
            hdr_cells[i].paragraphs[0].runs[0].bold = True

        for _, row in tmp.head(50).iterrows():
            cells = table.add_row().cells
            cells[0].text = str(row["source"])
            cells[1].text = str(row["published"])
            cells[2].text = str(row["title"])
            cells[3].text = str(row["relevance"])
            cells[4].text = str(row["impact"])
    doc.add_paragraph().add_run(" ").add_break()

    # --- Per-article blocks ---
    if results_df is not None and not results_df.empty:
        h = doc.add_paragraph("Articles & Auto-Summaries")
        _set_heading_style(h, size=14, bold=True)

        for i, row in results_df.iterrows():
            doc.add_paragraph().add_run(" ").add_break()
            p = doc.add_paragraph(f"{i+1}. {row['title']}")
            _set_heading_style(p, size=12, bold=True)
            
            meta = doc.add_paragraph()
            meta.add_run("Source: ").bold = True
            meta.add_run(str(row["source"]))
            meta.add_run("  |  Published: ").bold = True
            meta.add_run(str(row["published"]))
            meta.add_run("  |  Relevance: ").bold = True
            try:
                meta.add_run(f"{row['relevance']:.0%}")
            except Exception:
                meta.add_run(str(row["relevance"]))

            imp = ", ".join(row["impact"]) if isinstance(row["impact"], list) else str(row["impact"])
            impp = doc.add_paragraph()
            impp.add_run("Impact: ").bold = True
            impp.add_run(imp)

            sump = doc.add_paragraph()
            sump.add_run("Summary: ").bold = True
            sump.add_run(row.get("auto_summary") or "")

            linkp = doc.add_paragraph()
            linkp.add_run("Link: ").bold = True
            linkp.add_run(row.get("link") or "")

    # --- Social Sentiment ---
    if sent_summary and isinstance(sent_summary, dict) and (sent_summary.get("n", 0) > 0):
        doc.add_paragraph().add_run(" ").add_break()
        h2 = doc.add_paragraph("Twitter/X Sentiment Snapshot")
        _set_heading_style(h2, size=14, bold=True)
        _add_kv_para(doc, "Tweets analyzed", str(int(sent_summary.get("n", 0))))
        _add_kv_para(doc, "Mean (compound)", f"{sent_summary.get('mean_compound', 0.0):+.3f}")
        _add_kv_para(doc, "Positive", f"{100*sent_summary.get('share_pos',0.0):.1f}%")
        _add_kv_para(doc, "Neutral", f"{100*sent_summary.get('share_neu',0.0):.1f}%")
        _add_kv_para(doc, "Negative", f"{100*sent_summary.get('share_neg',0.0):.1f}%")

    # --- Daily Digest (Markdown) ---
    if digest_md:
        doc.add_paragraph().add_run(" ").add_break()
        h3 = doc.add_paragraph("Daily Digest")
        _set_heading_style(h3, size=14, bold=True)

        lines = (digest_md or "").splitlines()
        for line in lines:
            line = line.rstrip()
            if line.startswith("#"):
                txt = line.lstrip("# ").strip()
                p = doc.add_paragraph(txt)
                _set_heading_style(p, size=12, bold=True)
            elif line.startswith(("- ", "* ")):
                p = doc.add_paragraph(line[2:], style='List Bullet')
            elif line.strip() == "---":
                doc.add_paragraph().add_run("—" * 20)
            else:
                if line.strip():
                    doc.add_paragraph(line)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

# --- Markdown Digest Builder ---
def make_digest(df: pd.DataFrame, top_k: int = 12) -> str:
    header = f"# {APP_NAME} — Daily Digest\n\n*{TAGLINE}*\n\n> {QUOTE}\n\n"
    if df.empty:
        return header + "_No relevant items found for the selected period._"
    parts = [header, f"**Date:** {dt.datetime.now(dt.timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}\n\n"]
    for _, row in df.head(top_k).iterrows():
        impact = ", ".join(row["impact"])
        parts.append(
            f"### {row['title']}\n"
            f"- **Source:** {row['source']}  \n"
            f"- **Published:** {row['published']}  \n"
            f"- **Relevance:** {row['relevance']:.3f}  \n"
            f"- **Impact:** {impact}  \n"
            f"- **Summary:** {row['auto_summary']}\n"
            f"[Read more]({row['link']})\n\n---\n"
        )
    return "\n".join(parts)


# ==============================================================================
# 9. UI COMPONENTS (RENDERERS)
# ==============================================================================

def render_card(row: pd.Series):
    """Renders a single article card."""
    key = f"card_{hash_key(row['title'], row['link'])}"
    if "ai_analyses" not in st.session_state:
        st.session_state.ai_analyses = {}

    pub = row["published"]
    src = row["source"]
    rel = f"{row['relevance']:.0%}"
    title = row["title"]
    link = row["link"]
    img = row["image"] or FALLBACK_IMG
    summary = row["auto_summary"] or ""
    tags = row["impact"] or ["General"]

    with st.container():
        st.markdown(f"""
        <div class="card">
            <img class="thumb" src="{img}" alt="thumbnail">
            <div class="card-body">
                <div class="meta">{src} · {pub} · Relevance {rel}</div>
                <div class="title">{title}</div>
                <div class="badges">
                    {"".join([f'<span class="badge">{t}</span>' for t in tags])}
                </div>
                <div class="summary">{summary}</div>
                <a class="link" href="{link}" target="_blank">Read full article →</a>
            </div>
        </div>
        """, unsafe_allow_html=True)

        with st.expander("🔎 Analyze with AI", expanded=False):
            if not have_openai():
                st.warning("Add an `OPENAI_API_KEY` to your `.env` or Streamlit Secrets to run AI analysis.")
                st.info("Tip: open the 🧪 Diagnostics tab.")
            else:
                prev = st.session_state.ai_analyses.get(key)
                if prev:
                    st.markdown(prev)

                if st.button("Run LLM Analysis", key=f"btn_{key}", help="Run a deep-dive analysis on this article"):
                    with st.spinner("Analyzing article with AI..."):
                        full_text, _ = fetch_article_text_and_image(link)
                        body = full_text if len(full_text) > len(summary) else summary
                        if not body:
                            st.error("Could not extract article text to analyze.")
                        else:
                            md = analyze_with_llm(title, body, tags)
                            if not md:
                                st.error("AI analysis failed. Check Diagnostics and try again.")
                            else:
                                st.session_state["ai_analyses"][key] = md
                                st.markdown(md)
                                st.rerun() # Rerun to show the result without the button

# ==============================================================================
# 10. MAIN APP EXECUTION
# ==============================================================================

# --- Page Config & Initial State ---
st.set_page_config(page_title=APP_NAME, page_icon="🌍", layout="wide", initial_sidebar_state="expanded")
st.markdown(CARD_CSS, unsafe_allow_html=True)

# --- Durable stores (persist across reruns) ----
ss_get("results_df", None)
ss_get("results_digest_md", "")
ss_get("ai_analyses", {})
ss_get("last_scan_params", {})
ss_get("filters_impact", [])
ss_get("filters_source", [])
ss_get("sent_df", None)
ss_get("sent_summary", {})
init_chat_state() # Ensure chat history is set up

# ========================= UI: Sidebar =========================
with st.sidebar:
    st.title(f"🌍 {APP_NAME}")
    st.markdown(f"*{TAGLINE}*")
    st.markdown("---")

    # --- Primary Actions ---
    run_btn = st.button("🚀 Scan Now", use_container_width=True, type="primary")
    if st.button("♻️ Reset", use_container_width=True):
        for k in list(st.session_state.keys()):
            if k.startswith(("results_", "ai_", "chat_", "cfg_", "last_scan_", "filters_", "sent_")):
                del st.session_state[k]
        st.rerun()

    st.markdown("---")

    # --- Core Filters ---
    st.subheader("📅 Date Range")
    mode = st.radio("Mode", ["Quick Select", "Custom"], horizontal=True, key="date_mode", index=1)
    if mode == "Quick Select":
        quick = {"Last 24 Hours": 1, "Last 3 Days": 3, "Last Week": 7, "Last 2 Weeks": 14, "Last Month": 30}
        sel = st.selectbox("Window", list(quick.keys()), index=2, key="date_win")
        days = quick[sel]
        end_date = dt.datetime.now(dt.timezone.utc)
        start_date = end_date - dt.timedelta(days=days)
    else:
        d1, d2 = st.columns(2)
        with d1: sd = st.date_input("Start", value=dt.date.today() - dt.timedelta(days=7), key="start_date")
        with d2: ed = st.date_input("End", value=dt.date.today(), key="end_date")
        start_date = dt.datetime.combine(sd, dt.time.min, tzinfo=dt.timezone.utc)
        end_date = dt.datetime.combine(ed, dt.time.max, tzinfo=dt.timezone.utc)
    
    st.subheader("🎯 Relevance")
    min_relevance = st.slider("Min relevance (0.00–1.00)", 0.0, 1.0, 0.05, 0.01, "%.2f", key="min_rel")
    
    st.markdown("---")

    # --- Advanced Configuration ---
    with st.expander("⚙️ Advanced Configuration", expanded=False):
        # 📰 RSS/Atom Sources
        st.subheader("📰 RSS/Atom Sources")
        chosen_sources: List[str] = []
        for name, url in DEFAULT_SOURCES.items():
            if st.checkbox(name, value=True, key=f"src_{name}"):
                chosen_sources.append(url)
        
        st.subheader("➕ Custom RSS/Atom Sources")
        custom_text = st.text_area("Enter URLs (one per line)", height=80, key="custom_sources") or ""
        custom_urls = [u.strip() for u in custom_text.splitlines() if u.strip() and u.startswith(('http://', 'https://'))]
        if custom_text and not custom_urls:
            st.warning("No valid HTTP/HTTPS URLs found in custom sources.")
        for idx, url in enumerate(custom_urls):
            domain = urlparse(url).netloc or f"Custom {idx+1}"
            if st.checkbox(domain, value=True, key=f"custom_src_{idx}"):
                chosen_sources.append(url)
        
        if st.button("🔄 Check Feeds", key="check_feeds"):
            to_check = {name: url for name, url in DEFAULT_SOURCES.items()}
            for idx, url in enumerate(custom_urls):
                domain = urlparse(url).netloc or f"Custom {idx+1}"
                to_check[domain] = url
            for name, url in to_check.items():
                ok, status = validate_feed(url, ignore_recency_check=True)
                st.write(f"{'✅' if ok else '❌'} {name}: {status}")

        st.markdown("---")

        # 🧩 Newsdata.io (optional)
        st.subheader("🧩 Newsdata.io (optional)")
        use_newsdata = st.checkbox("Use Newsdata.io", value=True, key="use_nd")
        auto_key = get_newsdata_api_key()
        newsdata_key = st.text_input(
            "Newsdata.io API Key", 
            value=auto_key, 
            type="password", 
            key="nd_key_input",
            help="Loaded from NEWSDATA_API_KEY secret."
        )
        if use_newsdata:
            if newsdata_key:
                st.success("Using secured API key.")
            else:
                st.warning("No API key found. Add NEWSDATA_API_KEY to `.env`/Secrets.")

        newsdata_query = st.text_input("Query", value="tree crop commodities", key="nd_query")
        c1, c2, c3 = st.columns(3)
        with c1: nd_language = st.text_input("Language (e.g., en, fr)", value="", key="nd_lang")
        with c2: nd_country = st.text_input("Country (e.g., gh, ng, ci)", value="", key="nd_cty")
        with c3: nd_category = st.text_input("Category (e.g., business)", value="", key="nd_cat")
        nd_pages = st.number_input("Newsdata pages", 1, 10, 2, 1, key="nd_pages")

        st.markdown("---")

        # 🔍 Keywords
        st.subheader("🔍 Keywords")
        custom_kw = st.text_area("Keywords (comma-separated)", ", ".join(DEFAULT_KEYWORDS), height=100, key="kw_text")
        keywords = [k.strip() for k in custom_kw.split(",") if k.strip()]
        per_source_cap = st.number_input("Max articles per source", 1, 200, 10, 1, key="cap")

        st.markdown("---")

        # 📝 Content Settings
        st.subheader("📝 Content Settings")
        n_sent = st.number_input("Sentences per summary", 2, 10, 3, 1, key="n_sent")
        summary_method = st.radio("Summary Method", ["Extractive (Fast, Local)", "Abstractive (LLM)"], key="summary_method")
        if summary_method == "Abstractive (LLM)" and not have_openai():
            st.warning("No OPENAI_API_KEY found. LLM summaries will fall back to extractive method.")
        top_k = st.number_input("Digest: top items", 5, 100, 12, 1, key="top_k")

        st.markdown("---")

        # 🐦 Social Sentiment (Twitter/X)
        st.subheader("🐦 Social Sentiment (Twitter/X)")
        enable_social = st.checkbox("Enable Twitter/X sentiment", value=True, key="enable_social")
        sentiment_method = st.radio("Sentiment Method", ["Grok (xAI API)", "VADER (Local)"], key="sentiment_method")
        default_query = " OR ".join([kw for kw in keywords if " " not in kw][:6]) or "cashew OR shea OR cocoa"
        tw_query = st.text_input("Twitter search query", value=default_query, help="Example: cashew OR shea OR cocoa", key="tw_query")
        col_t1, col_t2, col_t3 = st.columns(3)
        with col_t1: tw_hours = st.number_input("Lookback hours", 6, 720, 72, 6, key="tw_hours")
        with col_t2: tw_lang = st.text_input("Language filter (e.g., en, fr)", value="en", key="tw_lang")
        with col_t3: tw_max = st.number_input("Max tweets", 10, 1000, 300, 50, key="tw_max")
        st.caption("Tip: 'Grok' uses XAI_API_KEY. 'VADER' uses TWITTER_BEARER_TOKEN for fetching.")

        st.markdown("---")

        # 🛡️ Resilience
        st.subheader("🛡️ Resilience")
        force_fetch = st.checkbox("⚡ Force RSS fetch if validation fails", value=True, key="force")
        ignore_recency = st.checkbox("🕒 Ignore RSS recency check", value=True, key="ignore_recent")
        dedupe_across_sources = st.checkbox("🧹 Deduplicate across sources", value=True, key="dedupe")

# --- Build Config Dict ---
current_params = {
    "chosen_sources": chosen_sources[:],
    "use_newsdata": bool(use_newsdata),
    "newsdata_key": newsdata_key,
    "newsdata_query": newsdata_query,
    "nd_language": nd_language,
    "nd_country": nd_country,
    "nd_category": nd_category,
    "nd_pages": int(nd_pages),
    "start_date": start_date,
    "end_date": end_date,
    "keywords": keywords[:],
    "min_relevance": float(min_relevance),
    "per_source_cap": int(per_source_cap),
    "n_sent": int(n_sent),
    "summary_method": summary_method,
    "top_k": int(top_k),
    "force_fetch": bool(force_fetch),
    "ignore_recency": bool(ignore_recency),
    "dedupe": bool(dedupe_across_sources),
    "enable_social": bool(enable_social),
    "sentiment_method": sentiment_method,
    "tw_query": tw_query,
    "tw_hours": int(tw_hours),
    "tw_lang": tw_lang.strip(),
    "tw_max": int(tw_max),
}


# ========================= UI: Main Page =========================

# --- Hero Banner ---
with st.container():
    st.markdown(f"""
    <div class="hero">
        <div class="pill">🌍 {APP_NAME}</div>
        <h1>{TAGLINE}</h1>
        <p>{QUOTE}</p>
    </div>
    """, unsafe_allow_html=True)
st.markdown("<br>", unsafe_allow_html=True)

# --- Main Tabbed Interface ---
tab1, tab2, tab3, tab4 = st.tabs([
    "📊 Pulse Dashboard",
    "🤖 Chat Assistant",
    "🔗 Quick Analyze URL",
    "🧪 Diagnostics"
])

# ========================= TAB 1: Pulse Dashboard =========================
with tab1:
    df = ss_get("results_df", None)
    sent_df = ss_get("sent_df", None)
    sent_summary = ss_get("sent_summary", {})

    if df is None:
        st.info("""
        **👋 Welcome to the Pulse Dashboard!**

        Click the **🚀 Scan Now** button in the sidebar to fetch, analyze, and display the latest market intelligence.
        
        **What this app does:**
        - 📰 Scans curated RSS feeds and optional Newsdata.io API.
        - 🎯 Scores relevance against your keywords (see ⚙️ Config).
        - 📝 Auto-summarizes (local-extractive or **LLM-abstractive**).
        - 🏷️ Tags each item (Supply Risk, FX & Policy, Logistics, etc.).
        - 🐦 (Optional) Collects and analyzes **Twitter/X sentiment**.
        - 💾 Outputs a **downloadable CSV**, **Word Report (.docx)** and **Markdown Digest**.
        """)
    else:
        # --- KPI Metrics ---
        st.subheader("Scan Snapshot")
        met1, met2, met3, met4 = st.columns(4)
        with met1:
            st.metric("Relevant Articles", len(df))
        with met2:
            st.metric("Avg. Relevance", f"{df['relevance'].mean():.1%}" if not df.empty else "N/A")
        with met3:
            st.metric("Tweets Analyzed", sent_summary.get("n", 0))
        with met4:
            st.metric("Mean Sentiment", f"{sent_summary.get('mean_compound', 0.0):+.3f}")

        st.markdown("---")

        # --- Filters ---
        st.subheader("Filter Results")
        c1, c2 = st.columns(2)
        with c1:
            all_impacts = sorted({t for tags in df["impact"] for t in tags})
            impact_filter = st.multiselect(
                "Filter by impact",
                options=all_impacts,
                default=ss_get("filters_impact", []),
                key="filters_impact"
            )
        with c2:
            source_filter = st.multiselect(
                "Filter by source",
                options=sorted(df["source"].unique()),
                default=ss_get("filters_source", []),
                key="filters_source"
            )

        filtered = df.copy()
        if impact_filter:
            filtered = filtered[filtered["impact"].apply(lambda x: any(t in x for t in impact_filter))]
        if source_filter:
            filtered = filtered[filtered["source"].isin(source_filter)]

        if filtered.empty and not df.empty:
            st.warning("No articles match your filters.")
        elif df.empty:
            st.warning("No relevant articles found. Try widening the date range or lowering the relevance threshold.")
            
        # --- Card Grid ---
        cards = list(filtered.to_dict("records"))
        n_cols = 3
        for i in range(0, len(cards), n_cols):
            cols = st.columns(n_cols)
            for j, col in enumerate(cols):
                if i + j < len(cards):
                    with col:
                        render_card(pd.Series(cards[i + j]))
        
        st.markdown("<br>", unsafe_allow_html=True)

        # --- Downloads & Digest Expander ---
        with st.expander("📝 View Digest & Download Reports", expanded=False):
            st.subheader("📝 Daily Digest")
            digest_md = make_digest(filtered if (impact_filter or source_filter) else df, top_k=current_params["top_k"])
            st.markdown(digest_md)

            st.subheader("⬇️ Downloads")
            ts = dt.datetime.now(dt.timezone.utc).strftime("%Y%m%d_%H%M%S")
            export_df = filtered if (impact_filter or source_filter) else df
            csv_name = f"oneafrica_pulse_{ts}.csv"
            md_name = f"oneafrica_pulse_digest_{ts}.md"
            st.download_button("📥 Download CSV", data=export_df.to_csv(index=False).encode("utf-8"),
                                file_name=csv_name, mime="text/csv")
            st.download_button("📥 Download Digest (Markdown)", data=digest_md.encode("utf-8"),
                                file_name=md_name, mime="text/markdown")

            if HAS_DOCX:
                try:
                    docx_bytes = build_results_docx(
                        app_name=APP_NAME,
                        tagline=TAGLINE,
                        quote=QUOTE,
                        results_df=export_df,
                        digest_md=digest_md,
                        params=ss_get("last_scan_params", current_params),
                        sent_summary=sent_summary
                    )
                    docx_name = f"oneafrica_pulse_{ts}.docx"
                    st.download_button(
                        "📥 Download Report (Word .docx)",
                        data=docx_bytes,
                        file_name=docx_name,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                except Exception as e:
                    st.error(f"Could not generate Word report: {e}")
            else:
                st.warning("Install `python-docx` to enable .docx downloads.")

        # --- Social Sentiment Section ---
        st.markdown("---")
        st.subheader("🐦 Social Sentiment — Twitter/X")
        if (sent_df is None) or (sent_df is not None and sent_df.empty):
            st.caption("No tweets captured for the current window/query. Enable in ⚙️ Config and click **Scan Now**.")
        else:
            share_df = pd.DataFrame({
                "label": ["Positive", "Neutral", "Negative"],
                "share": [
                    100 * sent_summary.get('share_pos', 0.0),
                    100 * sent_summary.get('share_neu', 0.0),
                    100 * sent_summary.get('share_neg', 0.0),
                ]
            })
            st.bar_chart(data=share_df, x="label", y="share", use_container_width=True)

            with st.expander("See recent tweets"):
                show_cols = ["created_at", "label", "compound", "text", "likes", "retweets", "username", "url"]
                st.dataframe(sent_df[show_cols], use_container_width=True, height=320)

# ========================= TAB 2: Chat Assistant =========================
with tab2:
    st.subheader("🤖 Chat Assistant")
    st.caption("Ask follow-ups, draft digests, or generate summaries based on the scan results.")

    # Render prior chat (omit system)
    for m in st.session_state.chat_history:
        if m["role"] == "system":
            continue
        with st.chat_message(m["role"]):
            st.markdown(m["content"])

    # Chat input
    user_prompt = st.chat_input("Type your message...")
    if user_prompt:
        st.session_state.chat_history.append({"role": "user", "content": user_prompt})
        with st.chat_message("user"):
            st.markdown(user_prompt)

        if not have_openai():
            with st.chat_message("assistant"):
                st.warning("No `OPENAI_API_KEY` found. Add it in ⚙️ Config or 🧪 Diagnostics.")
        else:
            reply, _streamed = generate_assistant_reply(st.session_state.chat_history)
            if reply:
                if not _streamed: # If streaming failed, render manually
                    with st.chat_message("assistant"):
                        st.markdown(reply)
                st.session_state.chat_history.append({"role": "assistant", "content": reply})
            else:
                with st.chat_message("assistant"):
                    st.error("The assistant is temporarily unavailable. Please try again in a moment.")

# ========================= TAB 3: Quick Analyze URL =========================
with tab3:
    st.subheader("🔗 Quick Analyze by URL (LLM)")
    st.caption("Paste any article URL for an instant AI-powered analysis.")
    
    quick_url = st.text_input("Paste any article URL", value="", placeholder="https://example.com/article")
    run_quick = st.button("Analyze URL", use_container_width=True, key="an_quick")

    if run_quick:
        if not have_openai():
            st.warning("Add an `OPENAI_API_KEY` to use AI analysis.")
        elif not quick_url:
            st.info("Please paste a valid URL.")
        else:
            with st.spinner("Fetching and analyzing..."):
                text, img = fetch_article_text_and_image(quick_url)
                if not text:
                    st.error("Could not extract article text from this URL.")
                else:
                    title_guess = text.split(".")[0][:140] if text else quick_url
                    tags = classify_impact(text)
                    md = analyze_with_llm(title_guess, text, tags)
                    if not md:
                        st.error("AI analysis failed. Please try again.")
                    else:
                        st.image(img, use_column_width=True)
                        st.markdown(f"**Source:** {urlparse(quick_url).netloc}")
                        st.markdown(md)

# ========================= TAB 4: Diagnostics =========================
with tab4:
    st.subheader("🧪 Diagnostics")
    st.caption("Check your environment and API connectivity.")

    st.markdown("**Core Libraries**")
    st.write(f"OPENAI package installed: **{OPENAI_OK}**")
    st.write(f"Scikit-learn (SK) package installed: **{HAS_SK}**")
    st.write(f"NLTK VADER package installed: **{HAS_VADER}**")
    st.write(f"Python-DOCX package installed: **{HAS_DOCX}**")
    st.markdown("---")
    
    st.markdown("**API Keys (from .env or st.secrets)**")
    st.write(f"OPENAI_API_KEY present: **{'Yes' if get_openai_api_key() else 'No'}**")
    st.write(f"OPENAI_BASE_URL: **{os.environ.get('OPENAI_BASE_URL','(not set)')}**")
    st.write(f"XAI_API_KEY (Grok) present: **{'Yes' if get_xai_api_key() else 'No'}**")
    st.write(f"TWITTER_BEARER_TOKEN (v2) present: **{'Yes' if get_twitter_bearer() else 'No'}**")
    st.write(f"NEWSDATA_API_KEY present: **{'Yes' if get_newsdata_api_key() else 'No'}**")
    st.markdown("---")

    st.markdown("**Connectivity Tests**")
    if st.button("Run OpenAI API Test"):
        if not have_openai():
            st.error("No OPENAI_API_KEY or package not installed.")
        else:
            client = get_openai_client()
            if client is None:
                st.error("Could not initialize OpenAI client (see logs).")
            else:
                try:
                    with st.spinner("Pinging gpt-4o-mini..."):
                        resp = client.chat.completions.create(
                            model="gpt-4o-mini", temperature=0,
                            messages=[{"role":"system","content":"You are a tester."},
                                      {"role":"user","content":"Reply with the single word: OK"}]
                        )
                        msg = (resp.choices[0].message.content or "").strip()
                        st.success(f"OpenAI replied: **{msg}**")
                except Exception as e:
                    st.error(f"Model call failed: {e}")

    if st.button("Run Grok API Test"):
        if not have_grok():
            st.error("No XAI_API_KEY or openai package not installed.")
        else:
            client = get_grok_client()
            if client is None:
                st.error("Could not initialize Grok client (see logs).")
            else:
                try:
                    with st.spinner("Pinging grok-beta..."):
                        resp = client.chat.completions.create(
                            model="grok-beta",
                            messages=[{"role": "user", "content": "Say hello!"}],
                            temperature=0.0,
                            max_tokens=10
                        )
                        msg = (resp.choices[0].message.content or "").strip()
                        st.success(f"Grok replied: **{msg}**")
                except Exception as e:
                    st.error(f"Grok call failed: {e}")

# ========================= Main Scan Logic =========================
if run_btn:
    try:
        if not current_params["chosen_sources"] and not (current_params["use_newsdata"] and current_params["newsdata_key"]):
            st.error("Pick at least one RSS source or enable Newsdata.io (see ⚙️ Config).")
        else:
            # 1. Fetch News
            rows = fetch_all(current_params)
            df = process_rows(rows)
            ss_set("results_df", df)
            ss_set("results_digest_md", make_digest(df, top_k=current_params["top_k"]))
            ss_set("last_scan_params", current_params)
            ss_set("ai_analyses", {}) # Clear old AI analyses
            
            # 2. Fetch Social Sentiment
            if current_params["enable_social"]:
                tweets = []
                df_t = pd.DataFrame()
                
                if current_params["sentiment_method"] == "Grok (xAI API)":
                    with st.sidebar.spinner("Running Grok sentiment..."):
                        q = current_params["tw_query"]
                        lang = current_params["tw_lang"]
                        hours = current_params["tw_hours"]
                        max_t = current_params["tw_max"]
                        tweets = fetch_tweets_via_grok(q, lang, hours, max_t)
                        df_t = analyze_tweet_sentiment_grok(tweets)

                else: # VADER (Local)
                    bearer = get_twitter_bearer()
                    if not bearer:
                        soft_fail("TWITTER_BEARER_TOKEN missing for VADER fetch.", "Skipping Twitter v2 fetch.")
                    else:
                        with st.sidebar.spinner("Fetching tweets (v2)..."):
                            q = current_params["tw_query"]
                            lang = current_params["tw_lang"]
                            hours = current_params["tw_hours"]
                            max_t = current_params["tw_max"]
                            start_iso = (dt.datetime.now(dt.timezone.utc) - dt.timedelta(hours=hours)).isoformat(timespec="seconds") + "Z"
                            tweets = fetch_tweets_via_twitter_v2(bearer, q, lang, start_iso, max_t)
                    
                    if tweets:
                        with st.sidebar.spinner("Analyzing with VADER..."):
                            df_t = analyze_tweet_sentiment(tweets)

                summ = summarize_sentiment(df_t)
                ss_set("sent_df", df_t)
                ss_set("sent_summary", summ)
            
            else: # Social disabled
                ss_set("sent_df", pd.DataFrame())
                ss_set("sent_summary", {"n": 0, "mean_compound": 0.0, "share_pos": 0.0, "share_neu": 1.0, "share_neg": 0.0})
            
            st.toast(f"Scan complete! Found {len(df)} articles.", icon="🎉")
            st.rerun() # Rerun to show new results in Tab 1

    except Exception as e:
        soft_fail("Something went wrong while assembling the results.", f"MAIN EXC {e}")
        st.error("We ran into a hiccup assembling the results. Please try again or adjust your filters.")
        logger.error(f"Main run_btn loop failed: {e}", exc_info=True)
    finally:
        # Display any non-fatal errors
        friendly_error_summary()